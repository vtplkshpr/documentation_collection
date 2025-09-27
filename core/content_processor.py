"""
Content processor for downloaded documents
"""
import logging
import os
from typing import Dict, Any, Optional
from pathlib import Path
import PyPDF2
from docx import Document
import pandas as pd
from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)

class ContentProcessor:
    """Content processor for various document types"""
    
    def __init__(self):
        self.supported_types = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_doc,
            '.html': self._process_html,
            '.htm': self._process_html,
            '.txt': self._process_txt,
            '.csv': self._process_csv,
            '.xlsx': self._process_excel,
            '.xls': self._process_excel
        }
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        Process document and extract metadata
        
        Args:
            file_path: Path to the document
        
        Returns:
            Dictionary with extracted metadata
        """
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return {}
        
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext not in self.supported_types:
            logger.warning(f"Unsupported file type: {file_ext}")
            return self._get_basic_metadata(file_path)
        
        try:
            processor = self.supported_types[file_ext]
            metadata = processor(file_path)
            metadata.update(self._get_basic_metadata(file_path))
            return metadata
        except Exception as e:
            logger.error(f"Error processing {file_path}: {e}")
            return self._get_basic_metadata(file_path)
    
    def _get_basic_metadata(self, file_path: str) -> Dict[str, Any]:
        """Get basic file metadata"""
        try:
            stat = os.stat(file_path)
            return {
                'file_size': stat.st_size,
                'file_name': os.path.basename(file_path),
                'file_path': file_path,
                'file_type': Path(file_path).suffix.lower(),
                'created_time': stat.st_ctime,
                'modified_time': stat.st_mtime
            }
        except Exception as e:
            logger.error(f"Error getting basic metadata for {file_path}: {e}")
            return {}
    
    def _process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Process PDF document"""
        metadata = {}
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Basic info
                metadata['page_count'] = len(pdf_reader.pages)
                metadata['file_type'] = 'PDF'
                
                # Extract text from first few pages
                text_content = ""
                max_pages = min(3, len(pdf_reader.pages))  # First 3 pages
                for i in range(max_pages):
                    page = pdf_reader.pages[i]
                    text_content += page.extract_text() + "\n"
                
                metadata['text_content'] = text_content[:1000]  # First 1000 chars
                metadata['has_text'] = len(text_content.strip()) > 0
                
                # PDF metadata
                if pdf_reader.metadata:
                    metadata['title'] = pdf_reader.metadata.get('/Title', '')
                    metadata['author'] = pdf_reader.metadata.get('/Author', '')
                    metadata['subject'] = pdf_reader.metadata.get('/Subject', '')
                    metadata['creator'] = pdf_reader.metadata.get('/Creator', '')
                
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
        
        return metadata
    
    def _process_docx(self, file_path: str) -> Dict[str, Any]:
        """Process DOCX document"""
        metadata = {}
        try:
            doc = Document(file_path)
            
            # Basic info
            metadata['file_type'] = 'DOCX'
            metadata['paragraph_count'] = len(doc.paragraphs)
            
            # Extract text
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            metadata['text_content'] = text_content[:1000]  # First 1000 chars
            metadata['has_text'] = len(text_content.strip()) > 0
            
            # Document properties
            if doc.core_properties:
                metadata['title'] = doc.core_properties.title or ''
                metadata['author'] = doc.core_properties.author or ''
                metadata['subject'] = doc.core_properties.subject or ''
                metadata['created'] = str(doc.core_properties.created) if doc.core_properties.created else ''
                metadata['modified'] = str(doc.core_properties.modified) if doc.core_properties.modified else ''
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
        
        return metadata
    
    def _process_doc(self, file_path: str) -> Dict[str, Any]:
        """Process DOC document (basic implementation)"""
        metadata = {'file_type': 'DOC'}
        # DOC processing requires additional libraries like python-docx2txt
        # For now, return basic metadata
        logger.warning(f"DOC processing not fully implemented for {file_path}")
        return metadata
    
    def _process_html(self, file_path: str) -> Dict[str, Any]:
        """Process HTML document"""
        metadata = {}
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
                soup = BeautifulSoup(content, 'html.parser')
                
                # Basic info
                metadata['file_type'] = 'HTML'
                metadata['title'] = soup.title.string if soup.title else ''
                
                # Extract text content
                text_content = soup.get_text()
                metadata['text_content'] = text_content[:1000]  # First 1000 chars
                metadata['has_text'] = len(text_content.strip()) > 0
                
                # Meta tags
                meta_tags = {}
                for meta in soup.find_all('meta'):
                    name = meta.get('name') or meta.get('property')
                    content = meta.get('content')
                    if name and content:
                        meta_tags[name] = content
                
                metadata['meta_tags'] = meta_tags
                
                # Links
                links = [a.get('href') for a in soup.find_all('a', href=True)]
                metadata['link_count'] = len(links)
                metadata['links'] = links[:10]  # First 10 links
                
        except Exception as e:
            logger.error(f"Error processing HTML {file_path}: {e}")
        
        return metadata
    
    def _process_txt(self, file_path: str) -> Dict[str, Any]:
        """Process text document"""
        metadata = {}
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
                
                metadata['file_type'] = 'TXT'
                metadata['text_content'] = content[:1000]  # First 1000 chars
                metadata['has_text'] = len(content.strip()) > 0
                metadata['line_count'] = len(content.splitlines())
                metadata['word_count'] = len(content.split())
                
        except Exception as e:
            logger.error(f"Error processing TXT {file_path}: {e}")
        
        return metadata
    
    def _process_csv(self, file_path: str) -> Dict[str, Any]:
        """Process CSV document"""
        metadata = {}
        try:
            df = pd.read_csv(file_path, nrows=5)  # First 5 rows
            
            metadata['file_type'] = 'CSV'
            metadata['columns'] = list(df.columns)
            metadata['row_count'] = len(df)
            metadata['column_count'] = len(df.columns)
            metadata['sample_data'] = df.head(3).to_dict('records')
            
        except Exception as e:
            logger.error(f"Error processing CSV {file_path}: {e}")
        
        return metadata
    
    def _process_excel(self, file_path: str) -> Dict[str, Any]:
        """Process Excel document"""
        metadata = {}
        try:
            # Read first sheet
            df = pd.read_excel(file_path, nrows=5)
            
            metadata['file_type'] = 'EXCEL'
            metadata['columns'] = list(df.columns)
            metadata['row_count'] = len(df)
            metadata['column_count'] = len(df.columns)
            metadata['sample_data'] = df.head(3).to_dict('records')
            
            # Get sheet names
            excel_file = pd.ExcelFile(file_path)
            metadata['sheet_names'] = excel_file.sheet_names
            
        except Exception as e:
            logger.error(f"Error processing Excel {file_path}: {e}")
        
        return metadata
    
    def extract_text_content(self, file_path: str) -> str:
        """
        Extract text content from document for AI analysis
        
        Args:
            file_path: Path to the document
        
        Returns:
            Extracted text content
        """
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return ""
        
        file_ext = Path(file_path).suffix.lower()
        
        try:
            if file_ext == '.pdf':
                return self._extract_pdf_text(file_path)
            elif file_ext in ['.html', '.htm']:
                return self._extract_html_text(file_path)
            elif file_ext == '.txt':
                return self._extract_txt_text(file_path)
            elif file_ext == '.docx':
                return self._extract_docx_text(file_path)
            else:
                logger.warning(f"Text extraction not supported for {file_ext}")
                return ""
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return ""
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_content = ""
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
                return text_content[:5000]  # Limit to 5000 chars for AI analysis
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return ""
    
    def _extract_html_text(self, file_path: str) -> str:
        """Extract text from HTML"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
                soup = BeautifulSoup(content, 'html.parser')
                text_content = soup.get_text()
                return text_content[:5000]  # Limit to 5000 chars for AI analysis
        except Exception as e:
            logger.error(f"Error extracting HTML text: {e}")
            return ""
    
    def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from TXT"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
                return content[:5000]  # Limit to 5000 chars for AI analysis
        except Exception as e:
            logger.error(f"Error extracting TXT text: {e}")
            return ""
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX"""
        try:
            doc = Document(file_path)
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            return text_content[:5000]  # Limit to 5000 chars for AI analysis
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            return ""

    def extract_keywords(self, text: str, max_keywords: int = 10) -> list:
        """
        Extract keywords from text (simple implementation)
        
        Args:
            text: Text to analyze
            max_keywords: Maximum number of keywords
        
        Returns:
            List of keywords
        """
        if not text:
            return []
        
        # Simple keyword extraction (in real implementation, use NLP libraries)
        words = text.lower().split()
        
        # Filter out common words
        stop_words = {'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by', 'is', 'are', 'was', 'were', 'be', 'been', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would', 'could', 'should'}
        
        filtered_words = [word for word in words if len(word) > 3 and word not in stop_words]
        
        # Count frequency
        word_count = {}
        for word in filtered_words:
            word_count[word] = word_count.get(word, 0) + 1
        
        # Sort by frequency
        sorted_words = sorted(word_count.items(), key=lambda x: x[1], reverse=True)
        
        return [word for word, count in sorted_words[:max_keywords]]
